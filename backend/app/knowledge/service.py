"""
Knowledge Base Service for RAG (Retrieval Augmented Generation).

This module provides:
- Document upload and processing (PDF, Word, Excel, CSV, URLs)
- Text extraction and chunking
- Embedding generation and storage
- Semantic similarity search for context retrieval
"""

import os
import re
import json
import hashlib
import asyncio
from typing import List, Dict, Optional, Any, Tuple
from datetime import datetime
from pathlib import Path
import httpx

from sqlalchemy.orm import Session
from sqlalchemy import desc, and_, or_

from app.core.config import settings
from app.core.logging import logger
from app.models import (
    KnowledgeDocument, KnowledgeChunk, 
    KnowledgeDocumentType, KnowledgeDocumentStatus, User
)


# Storage path for uploaded files - use relative path as fallback for local dev
_data_dir = getattr(settings, 'DATA_DIR', None) or os.environ.get('DATA_DIR', './data')
KNOWLEDGE_STORAGE_PATH = Path(_data_dir) / "knowledge"
try:
    KNOWLEDGE_STORAGE_PATH.mkdir(parents=True, exist_ok=True)
except OSError as e:
    # Fallback to local directory if system path fails
    KNOWLEDGE_STORAGE_PATH = Path("./data/knowledge")
    KNOWLEDGE_STORAGE_PATH.mkdir(parents=True, exist_ok=True)
    logger.warning(f"Using fallback knowledge storage path: {KNOWLEDGE_STORAGE_PATH}")

# Chunk configuration
DEFAULT_CHUNK_SIZE = 1000  # characters
DEFAULT_CHUNK_OVERLAP = 200  # characters


class DocumentProcessor:
    """Processes various document formats and extracts text."""
    
    @staticmethod
    async def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            import pypdf
            text_parts = []
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return "\n\n".join(text_parts)
        except ImportError:
            logger.warning("pypdf not installed, trying fallback")
            # Fallback: try pdfminer
            try:
                from pdfminer.high_level import extract_text
                return extract_text(file_path)
            except ImportError:
                raise ImportError("Install pypdf or pdfminer.six for PDF support: pip install pypdf")
    
    @staticmethod
    async def extract_from_docx(file_path: str) -> str:
        """Extract text from Word document."""
        try:
            from docx import Document
            doc = Document(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("Install python-docx for Word support: pip install python-docx")
    
    @staticmethod
    async def extract_from_excel(file_path: str) -> str:
        """Extract text from Excel file."""
        try:
            import pandas as pd
            text_parts = []
            
            # Read all sheets
            xlsx = pd.ExcelFile(file_path)
            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)
                text_parts.append(f"=== Sheet: {sheet_name} ===")
                # Convert to readable format
                for idx, row in df.iterrows():
                    row_text = " | ".join(f"{col}: {val}" for col, val in row.items() if pd.notna(val))
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
        except ImportError:
            raise ImportError("Install pandas and openpyxl for Excel support: pip install pandas openpyxl")
    
    @staticmethod
    async def extract_from_csv(file_path: str) -> str:
        """Extract text from CSV file."""
        try:
            import pandas as pd
            df = pd.read_csv(file_path)
            text_parts = []
            
            # Include column headers
            text_parts.append("Columns: " + ", ".join(df.columns.tolist()))
            
            # Convert rows
            for idx, row in df.iterrows():
                row_text = " | ".join(f"{col}: {val}" for col, val in row.items() if pd.notna(val))
                if row_text:
                    text_parts.append(row_text)
            
            return "\n".join(text_parts)
        except ImportError:
            raise ImportError("Install pandas for CSV support: pip install pandas")
    
    @staticmethod
    async def extract_from_url(url: str) -> str:
        """Extract text content from a URL."""
        try:
            async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
                response = await client.get(url)
                response.raise_for_status()
                content_type = response.headers.get('content-type', '')
                
                if 'text/html' in content_type:
                    # Parse HTML
                    try:
                        from bs4 import BeautifulSoup
                        soup = BeautifulSoup(response.text, 'html.parser')
                        
                        # Remove scripts and styles
                        for script in soup(["script", "style", "nav", "footer", "header"]):
                            script.decompose()
                        
                        # Get text
                        text = soup.get_text(separator='\n')
                        # Clean up whitespace
                        lines = [line.strip() for line in text.splitlines() if line.strip()]
                        return "\n".join(lines)
                    except ImportError:
                        # Fallback: basic text extraction
                        return response.text
                else:
                    # Plain text or other
                    return response.text
        except Exception as e:
            logger.error("url_extraction_failed", url=url, error=str(e))
            raise
    
    @staticmethod
    async def extract_from_text(file_path: str) -> str:
        """Extract text from plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    @classmethod
    async def extract(cls, source_type: str, source_path: str, mime_type: str = None) -> str:
        """Extract text from any supported source."""
        if source_type == "url":
            return await cls.extract_from_url(source_path)
        
        # File-based extraction
        if mime_type:
            if 'pdf' in mime_type:
                return await cls.extract_from_pdf(source_path)
            elif 'word' in mime_type or 'docx' in mime_type:
                return await cls.extract_from_docx(source_path)
            elif 'excel' in mime_type or 'spreadsheet' in mime_type:
                return await cls.extract_from_excel(source_path)
            elif 'csv' in mime_type:
                return await cls.extract_from_csv(source_path)
        
        # Fallback to extension-based detection
        ext = Path(source_path).suffix.lower()
        extractors = {
            '.pdf': cls.extract_from_pdf,
            '.docx': cls.extract_from_docx,
            '.doc': cls.extract_from_docx,
            '.xlsx': cls.extract_from_excel,
            '.xls': cls.extract_from_excel,
            '.csv': cls.extract_from_csv,
            '.txt': cls.extract_from_text,
            '.md': cls.extract_from_text,
            '.json': cls.extract_from_text,
        }
        
        extractor = extractors.get(ext, cls.extract_from_text)
        return await extractor(source_path)


class TextChunker:
    """Splits text into chunks for embedding."""
    
    @staticmethod
    def chunk_text(
        text: str, 
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP
    ) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks."""
        if not text:
            return []
        
        # Clean text
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings
                for sep in ['. ', '.\n', '!\n', '?\n', '\n\n']:
                    last_sep = text[start:end].rfind(sep)
                    if last_sep > chunk_size * 0.5:  # At least half the chunk
                        end = start + last_sep + len(sep)
                        break
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunks.append({
                    "index": chunk_index,
                    "content": chunk_text,
                    "start_char": start,
                    "end_char": end,
                    "token_count": len(chunk_text.split())  # Rough estimate
                })
                chunk_index += 1
            
            # Move start with overlap
            start = end - chunk_overlap if end < len(text) else len(text)
        
        return chunks


class EmbeddingService:
    """Generates embeddings for text chunks using available models."""
    
    def __init__(self):
        self.model_name = None
        self._embedding_dim = 384  # Default for sentence-transformers
    
    async def get_embedding(self, text: str) -> Optional[List[float]]:
        """Get embedding for text using Ollama or fallback."""
        # Try Ollama first
        try:
            return await self._get_ollama_embedding(text)
        except Exception as e:
            logger.warning("ollama_embedding_failed", error=str(e))
        
        # Fallback to simple TF-IDF style embedding
        return self._get_simple_embedding(text)
    
    async def _get_ollama_embedding(self, text: str) -> List[float]:
        """Get embedding from Ollama."""
        ollama_url = settings.OLLAMA_BASE_URL or "http://host.docker.internal:11434"
        
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                f"{ollama_url}/api/embeddings",
                json={
                    "model": settings.OLLAMA_MODEL or "llama3:latest",
                    "prompt": text[:2000]  # Limit context
                }
            )
            response.raise_for_status()
            data = response.json()
            self.model_name = f"ollama:{settings.OLLAMA_MODEL}"
            return data.get("embedding", [])
    
    def _get_simple_embedding(self, text: str) -> List[float]:
        """Create a simple hash-based embedding as fallback."""
        # This is a very basic fallback - not suitable for production
        # In production, use sentence-transformers or OpenAI embeddings
        import hashlib
        
        self.model_name = "simple_hash"
        
        # Create a deterministic embedding from text
        words = text.lower().split()[:100]
        embedding = []
        
        for i in range(self._embedding_dim):
            hash_input = f"{i}:{' '.join(words)}"
            hash_val = int(hashlib.md5(hash_input.encode()).hexdigest(), 16)
            # Normalize to -1 to 1
            embedding.append((hash_val % 1000) / 500 - 1)
        
        return embedding
    
    @staticmethod
    def cosine_similarity(a: List[float], b: List[float]) -> float:
        """Calculate cosine similarity between two embeddings."""
        if not a or not b or len(a) != len(b):
            return 0.0
        
        dot_product = sum(x * y for x, y in zip(a, b))
        norm_a = sum(x ** 2 for x in a) ** 0.5
        norm_b = sum(x ** 2 for x in b) ** 0.5
        
        if norm_a == 0 or norm_b == 0:
            return 0.0
        
        return dot_product / (norm_a * norm_b)


class KnowledgeService:
    """Main service for knowledge base management and retrieval."""
    
    def __init__(self, db: Session):
        self.db = db
        self.processor = DocumentProcessor()
        self.chunker = TextChunker()
        self.embedder = EmbeddingService()
    
    def compute_content_hash(self, content: str) -> str:
        """Compute SHA-256 hash of content for deduplication."""
        return hashlib.sha256(content.encode('utf-8')).hexdigest()
    
    def check_duplicate(
        self, 
        content_hash: str, 
        scope: str = "global",
        user_id: int = None
    ) -> Optional[KnowledgeDocument]:
        """Check if a document with the same hash already exists."""
        query = self.db.query(KnowledgeDocument).filter(
            KnowledgeDocument.content_hash == content_hash,
            KnowledgeDocument.is_active == True
        )
        
        # For user scope, also check global (admin) documents
        if scope == "user":
            # Check if exists in admin-managed KB
            admin_doc = query.filter(KnowledgeDocument.is_admin_managed == True).first()
            if admin_doc:
                return admin_doc
        
        return query.first()
    
    async def add_document(
        self,
        title: str,
        source_type: str,  # "file" or "url"
        user_id: int,
        file_path: str = None,
        file_name: str = None,
        source_url: str = None,
        doc_type: KnowledgeDocumentType = KnowledgeDocumentType.CUSTOM,
        description: str = None,
        target_functions: List[str] = None,
        target_platforms: List[str] = None,
        tags: List[str] = None,
        priority: int = 5,
        scope: str = "global",  # "global" (admin) or "user" (individual)
        is_admin_managed: bool = True,
        crawl_depth: int = 0
    ) -> KnowledgeDocument:
        """Add a new document to the knowledge base."""
        
        # Determine file info
        mime_type = None
        file_size = None
        content_hash = None
        
        if source_type == "file" and file_path:
            file_size = os.path.getsize(file_path)
            ext = Path(file_path).suffix.lower()
            mime_types = {
                '.pdf': 'application/pdf',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.doc': 'application/msword',
                '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                '.csv': 'text/csv',
                '.txt': 'text/plain',
                '.md': 'text/markdown',
            }
            mime_type = mime_types.get(ext, 'application/octet-stream')
            
            # Compute file hash for deduplication
            with open(file_path, 'rb') as f:
                content_hash = hashlib.sha256(f.read()).hexdigest()
        elif source_type == "url" and source_url:
            # For URLs, hash the URL itself (content hash will be computed after fetch)
            content_hash = hashlib.sha256(source_url.encode('utf-8')).hexdigest()
        
        # Check for duplicates
        if content_hash:
            existing = self.check_duplicate(content_hash, scope, user_id)
            if existing:
                if existing.is_admin_managed and not is_admin_managed:
                    raise ValueError(
                        f"This document already exists in the admin-managed knowledge base: '{existing.title}'"
                    )
                elif existing.uploaded_by_id == user_id:
                    raise ValueError(f"You have already uploaded this document: '{existing.title}'")
        
        # Create document record
        doc = KnowledgeDocument(
            title=title,
            description=description,
            doc_type=doc_type,
            scope=scope,
            is_admin_managed=is_admin_managed,
            source_type=source_type,
            source_url=source_url,
            file_name=file_name,
            file_path=file_path,
            file_size=file_size,
            mime_type=mime_type,
            content_hash=content_hash,
            crawl_depth=crawl_depth,
            status=KnowledgeDocumentStatus.PENDING,
            target_functions=target_functions,
            target_platforms=target_platforms,
            tags=tags,
            priority=priority if is_admin_managed else max(1, priority - 2),  # Admin docs get higher priority
            uploaded_by_id=user_id
        )
        
        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)
        
        logger.info("knowledge_document_added", 
                   doc_id=doc.id, 
                   title=title, 
                   user_id=user_id,
                   scope=scope,
                   is_admin_managed=is_admin_managed)
        
        return doc
    
    async def process_document(self, doc_id: int) -> KnowledgeDocument:
        """Process a document: extract text, chunk, and create embeddings."""
        doc = self.db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
        if not doc:
            raise ValueError(f"Document {doc_id} not found")
        
        try:
            doc.status = KnowledgeDocumentStatus.PROCESSING
            self.db.commit()
            
            # Extract text
            source_path = doc.source_url if doc.source_type == "url" else doc.file_path
            raw_content = await self.processor.extract(
                doc.source_type,
                source_path,
                doc.mime_type
            )
            
            doc.raw_content = raw_content
            
            # Chunk the text
            chunks = self.chunker.chunk_text(raw_content)
            doc.chunk_count = len(chunks)
            
            # Delete existing chunks
            self.db.query(KnowledgeChunk).filter(
                KnowledgeChunk.document_id == doc_id
            ).delete()
            
            # Create embeddings and store chunks
            for chunk_data in chunks:
                embedding = await self.embedder.get_embedding(chunk_data["content"])
                
                chunk = KnowledgeChunk(
                    document_id=doc_id,
                    chunk_index=chunk_data["index"],
                    content=chunk_data["content"],
                    token_count=chunk_data["token_count"],
                    embedding=embedding,
                    embedding_model=self.embedder.model_name,
                    chunk_metadata={
                        "start_char": chunk_data["start_char"],
                        "end_char": chunk_data["end_char"]
                    }
                )
                self.db.add(chunk)
            
            doc.status = KnowledgeDocumentStatus.READY
            self.db.commit()
            
            logger.info("knowledge_document_processed", 
                       doc_id=doc_id, 
                       chunks=len(chunks),
                       content_length=len(raw_content))
            
            return doc
            
        except Exception as e:
            doc.status = KnowledgeDocumentStatus.FAILED
            doc.processing_error = str(e)
            self.db.commit()
            logger.error("knowledge_document_processing_failed", doc_id=doc_id, error=str(e))
            raise
    
    async def search(
        self,
        query: str,
        target_function: str = None,
        target_platform: str = None,
        doc_type: KnowledgeDocumentType = None,
        top_k: int = 5,
        min_similarity: float = 0.3,
        uploaded_by_id: Optional[int] = None,
        include_admin_managed: bool = True,
        include_user_managed: bool = True,
    ) -> List[Dict[str, Any]]:
        """Search knowledge base for relevant chunks."""
        # Get query embedding
        query_embedding = await self.embedder.get_embedding(query)
        
        if not query_embedding:
            return []
        
        # Build document filter
        doc_query = self.db.query(KnowledgeDocument).filter(
            KnowledgeDocument.status == KnowledgeDocumentStatus.READY,
            KnowledgeDocument.is_active == True
        )

        if uploaded_by_id is not None:
            allowed = []
            if include_admin_managed:
                allowed.append(KnowledgeDocument.is_admin_managed == True)
            if include_user_managed:
                allowed.append(
                    and_(
                        KnowledgeDocument.is_admin_managed == False,
                        KnowledgeDocument.uploaded_by_id == uploaded_by_id,
                    )
                )
            if allowed:
                doc_query = doc_query.filter(or_(*allowed))
            else:
                return []
        
        if doc_type:
            doc_query = doc_query.filter(KnowledgeDocument.doc_type == doc_type)
        
        # Get all relevant chunks
        docs = doc_query.all()
        doc_ids = [d.id for d in docs]
        
        # Filter by target function/platform if specified
        if target_function or target_platform:
            filtered_ids = []
            for doc in docs:
                if target_function and doc.target_functions:
                    if target_function not in doc.target_functions:
                        continue
                if target_platform and doc.target_platforms:
                    if target_platform not in doc.target_platforms:
                        continue
                filtered_ids.append(doc.id)
            doc_ids = filtered_ids
        
        if not doc_ids:
            return []
        
        chunks = self.db.query(KnowledgeChunk).filter(
            KnowledgeChunk.document_id.in_(doc_ids)
        ).all()
        
        # Calculate similarities
        results = []
        for chunk in chunks:
            if not chunk.embedding:
                continue
            
            similarity = EmbeddingService.cosine_similarity(
                query_embedding, 
                chunk.embedding
            )
            
            if similarity >= min_similarity:
                results.append({
                    "chunk_id": chunk.id,
                    "document_id": chunk.document_id,
                    "document_title": chunk.document.title,
                    "doc_type": chunk.document.doc_type.value if chunk.document.doc_type else None,
                    "content": chunk.content,
                    "similarity": similarity,
                    "priority": chunk.document.priority,
                    "tags": chunk.document.tags
                })
        
        # Sort by similarity * priority
        results.sort(key=lambda x: x["similarity"] * (x["priority"] / 10), reverse=True)
        
        # Update usage stats
        for result in results[:top_k]:
            doc = self.db.query(KnowledgeDocument).filter(
                KnowledgeDocument.id == result["document_id"]
            ).first()
            if doc:
                doc.usage_count = (doc.usage_count or 0) + 1
                doc.last_used_at = datetime.utcnow()
        
        self.db.commit()
        
        return results[:top_k]
    
    def get_context_for_prompt(
        self,
        query: str,
        target_function: str = None,
        target_platform: str = None,
        max_tokens: int = 2000
    ) -> Tuple[str, List[Dict]]:
        """Get formatted context for prompt injection from RAG search.
        
        Returns:
            Tuple of (formatted_context_string, source_documents)
        """
        import asyncio
        
        # Run async search
        loop = asyncio.get_event_loop()
        if loop.is_running():
            # Create a new loop for nested async
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(
                    asyncio.run,
                    self.search(query, target_function, target_platform)
                )
                results = future.result()
        else:
            results = asyncio.run(self.search(query, target_function, target_platform))
        
        if not results:
            return "", []
        
        # Build context string
        context_parts = []
        sources = []
        current_tokens = 0
        
        for result in results:
            chunk_tokens = result.get("content", "").split().__len__()
            if current_tokens + chunk_tokens > max_tokens:
                break
            
            context_parts.append(f"""
=== From: {result['document_title']} ({result['doc_type']}) ===
{result['content']}
""")
            sources.append({
                "id": result["document_id"],
                "title": result["document_title"],
                "similarity": round(result["similarity"], 3)
            })
            current_tokens += chunk_tokens
        
        context_str = "\n".join(context_parts)
        
        return context_str, sources
    
    def list_documents(
        self,
        doc_type: KnowledgeDocumentType = None,
        status: KnowledgeDocumentStatus = None,
        is_active: bool = None,
        limit: int = 100
    ) -> List[KnowledgeDocument]:
        """List documents in the knowledge base."""
        query = self.db.query(KnowledgeDocument)
        
        if doc_type:
            query = query.filter(KnowledgeDocument.doc_type == doc_type)
        if status:
            query = query.filter(KnowledgeDocument.status == status)
        if is_active is not None:
            query = query.filter(KnowledgeDocument.is_active == is_active)
        
        return query.order_by(desc(KnowledgeDocument.created_at)).limit(limit).all()
    
    def delete_document(self, doc_id: int) -> bool:
        """Delete a document and its chunks using raw SQL to avoid ORM tracking issues."""
        from sqlalchemy import text
        from app.models import KnowledgeChunk
        
        # Query document metadata WITHOUT loading relationships
        doc_query = self.db.query(
            KnowledgeDocument.title, 
            KnowledgeDocument.file_path
        ).filter(KnowledgeDocument.id == doc_id).first()
        
        if not doc_query:
            return False
        
        doc_title = doc_query.title
        file_path = doc_query.file_path
        
        try:
            # Use raw SQL to delete chunks first
            chunk_result = self.db.execute(
                text("DELETE FROM knowledge_chunks WHERE document_id = :doc_id"),
                {"doc_id": doc_id}
            )
            chunk_count = chunk_result.rowcount
            
            # Use raw SQL to delete document
            doc_result = self.db.execute(
                text("DELETE FROM knowledge_documents WHERE id = :doc_id"),
                {"doc_id": doc_id}
            )
            doc_count = doc_result.rowcount
            
            # Commit the transaction
            self.db.commit()
            
            # Delete file if exists (after successful DB commit)
            if file_path and os.path.exists(file_path):
                try:
                    os.remove(file_path)
                    logger.info("file_deleted", path=file_path)
                except Exception as e:
                    logger.warning("failed_to_delete_file", path=file_path, error=str(e))
            
            logger.info("knowledge_document_deleted", 
                       doc_id=doc_id, 
                       title=doc_title, 
                       chunks_deleted=chunk_count,
                       doc_deleted=doc_count)
            return True
            
        except Exception as e:
            self.db.rollback()
            logger.error("failed_to_delete_document", doc_id=doc_id, error=str(e))
            return False
    
    def update_document(
        self,
        doc_id: int,
        title: str = None,
        description: str = None,
        target_functions: List[str] = None,
        target_platforms: List[str] = None,
        tags: List[str] = None,
        priority: int = None,
        is_active: bool = None
    ) -> Optional[KnowledgeDocument]:
        """Update document metadata."""
        doc = self.db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
        if not doc:
            return None
        
        if title is not None:
            doc.title = title
        if description is not None:
            doc.description = description
        if target_functions is not None:
            doc.target_functions = target_functions
        if target_platforms is not None:
            doc.target_platforms = target_platforms
        if tags is not None:
            doc.tags = tags
        if priority is not None:
            doc.priority = priority
        if is_active is not None:
            doc.is_active = is_active
        
        self.db.commit()
        self.db.refresh(doc)
        
        return doc


def get_knowledge_service(db: Session) -> KnowledgeService:
    """Factory function for KnowledgeService."""
    return KnowledgeService(db)
