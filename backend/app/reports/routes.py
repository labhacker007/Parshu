"""Report generation and sharing APIs."""
import csv
import io
from fastapi import APIRouter, Depends, HTTPException, status, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import desc, and_
from app.core.database import get_db
from app.auth.dependencies import get_current_user, require_permission
from app.auth.rbac import Permission
from app.models import Report, ReportStatus, Article, ArticleStatus, User, ExtractedIntelligence, Hunt, HuntExecution, AuditEventType
from app.core.logging import logger
from app.audit.manager import AuditManager
from pydantic import BaseModel
from typing import List, Optional, Literal
from datetime import datetime, timedelta


class ReportCreateRequest(BaseModel):
    title: str
    article_ids: List[int]
    report_type: str = "comprehensive"  # comprehensive, executive, technical


class AutoReportRequest(BaseModel):
    report_type: Literal["executive", "technical", "comprehensive"] = "executive"
    period: Literal["daily", "weekly"] = "daily"
    include_high_priority_only: bool = False


class ReportResponse(BaseModel):
    id: int
    title: str
    article_ids: List[int]
    content: Optional[str] = None
    executive_summary: Optional[str] = None
    technical_summary: Optional[str] = None
    key_findings: List[str] = []
    recommendations: List[str] = []
    report_type: str
    status: str = "DRAFT"
    generated_by_id: Optional[int] = None
    generated_at: datetime
    edited_by_id: Optional[int] = None
    edited_at: Optional[datetime] = None
    published_by_id: Optional[int] = None
    published_at: Optional[datetime] = None
    version: int = 1
    shared_with_emails: List[str]
    created_at: datetime
    updated_at: datetime
    
    class Config:
        from_attributes = True


class ReportUpdateRequest(BaseModel):
    """Request to update a draft report."""
    title: Optional[str] = None
    executive_summary: Optional[str] = None
    technical_summary: Optional[str] = None
    key_findings: Optional[List[str]] = None
    recommendations: Optional[List[str]] = None
    content: Optional[str] = None


class ReportPublishRequest(BaseModel):
    """Request to publish a report."""
    notes: Optional[str] = None  # Optional publishing notes


router = APIRouter(prefix="/reports", tags=["reports"])


class ReportCreateWithGenAIRequest(BaseModel):
    title: str
    article_ids: List[int]
    report_type: str = "comprehensive"  # comprehensive, executive, technical
    use_genai: bool = True  # Use GenAI for enhanced report generation


@router.post("/", response_model=ReportResponse)
async def create_report(
    request: ReportCreateWithGenAIRequest,
    current_user: User = Depends(require_permission(Permission.CREATE_REPORTS.value)),
    db: Session = Depends(get_db)
):
    """Generate a new report from selected articles.
    
    Set use_genai=True to use AI for enhanced summarization and synthesis.
    """
    # Fetch articles
    articles = db.query(Article).filter(Article.id.in_(request.article_ids)).all()
    if not articles:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="No articles found")
    
    # Generate content based on report type
    exec_summary = None
    tech_summary = None
    content = None
    key_findings = []
    recommendations = []
    
    if request.use_genai:
        # Use GenAI for enhanced reports
        if request.report_type == "executive":
            exec_summary = await _generate_executive_summary_genai(articles, db)
            content = exec_summary
        elif request.report_type == "technical":
            tech_summary = await _generate_technical_summary_genai(articles, db)
            content = tech_summary
        else:
            # Comprehensive: combine both
            exec_summary = await _generate_executive_summary_genai(articles, db)
            tech_summary = await _generate_technical_summary_genai(articles, db)
            content = f"# Executive Summary\n\n{exec_summary}\n\n---\n\n# Technical Details\n\n{tech_summary}"
        
        # Extract key findings and recommendations from articles
        key_findings = _extract_key_findings(articles)
        recommendations = _extract_recommendations(articles)
    else:
        # Fallback to simple concatenation
        if request.report_type == "executive":
            content = _generate_executive_summary(articles)
            exec_summary = content
        elif request.report_type == "technical":
            content = _generate_technical_summary(articles)
            tech_summary = content
        else:
            content = _generate_comprehensive_report(articles)
    
    # Create report in DRAFT status for review/editing
    from app.models import ReportStatus
    report = Report(
        title=request.title,
        article_ids=request.article_ids,
        content=content,
        executive_summary=exec_summary,
        technical_summary=tech_summary,
        key_findings=key_findings,
        recommendations=recommendations,
        report_type=request.report_type,
        status=ReportStatus.DRAFT,
        generated_by_id=current_user.id,
        generated_at=datetime.utcnow(),
        version=1
    )
    
    db.add(report)
    db.commit()
    db.refresh(report)
    
    method = "genai" if request.use_genai else "fallback"
    logger.info("report_generated", report_id=report.id, user_id=current_user.id, article_count=len(articles), method=method)
    
    # Audit log
    AuditManager.log_event(
        db=db,
        user_id=current_user.id,
        event_type=AuditEventType.REPORT_GENERATION,
        action=f"Generated {request.report_type} report: {request.title}",
        resource_type="report",
        resource_id=report.id,
        details={"articles_count": len(articles), "method": method}
    )
    
    return ReportResponse.model_validate(report)


@router.get("/{report_id}", response_model=ReportResponse)
def get_report(
    report_id: int,
    current_user: User = Depends(require_permission(Permission.READ_REPORTS.value)),
    db: Session = Depends(get_db)
):
    """Retrieve a report."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
    
    return ReportResponse.model_validate(report)


@router.patch("/{report_id}", response_model=ReportResponse)
def update_report(
    report_id: int,
    request: ReportUpdateRequest,
    current_user: User = Depends(require_permission(Permission.EDIT_REPORTS.value)),
    db: Session = Depends(get_db)
):
    """Update a draft report."""
    
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
    
    from app.models import ReportStatus
    
    # Can only edit DRAFT reports
    if report.status != ReportStatus.DRAFT:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Can only edit reports in DRAFT status. Published reports cannot be modified."
        )
    
    # Update fields
    update_data = request.model_dump(exclude_unset=True)
    for field, value in update_data.items():
        setattr(report, field, value)
    
    # Track edit metadata
    report.edited_by_id = current_user.id
    report.edited_at = datetime.utcnow()
    report.version += 1
    
    db.commit()
    db.refresh(report)
    
    logger.info("report_edited", report_id=report_id, editor_id=current_user.id, version=report.version)
    
    # Audit log
    AuditManager.log_event(
        db=db,
        user_id=current_user.id,
        event_type=AuditEventType.REPORT_GENERATION,
        action=f"Edited report (v{report.version}): {report.title}",
        resource_type="report",
        resource_id=report_id,
        details={"version": report.version, "fields_updated": list(update_data.keys())}
    )
    
    return ReportResponse.model_validate(report)


@router.post("/{report_id}/publish", response_model=ReportResponse)
def publish_report(
    report_id: int,
    request: ReportPublishRequest = ReportPublishRequest(),
    current_user: User = Depends(require_permission(Permission.PUBLISH_REPORTS.value)),
    db: Session = Depends(get_db)
):
    """Publish a report, making it final and ready for viewing/download."""
    
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
    
    from app.models import ReportStatus
    
    # Can only publish DRAFT reports
    if report.status == ReportStatus.PUBLISHED:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Report is already published"
        )
    
    # Update status to published
    report.status = ReportStatus.PUBLISHED
    report.published_by_id = current_user.id
    report.published_at = datetime.utcnow()
    
    db.commit()
    db.refresh(report)
    
    logger.info("report_published", report_id=report_id, publisher_id=current_user.id)
    
    # Audit log
    AuditManager.log_event(
        db=db,
        user_id=current_user.id,
        event_type=AuditEventType.REPORT_GENERATION,
        action=f"Published report: {report.title}",
        resource_type="report",
        resource_id=report_id,
        details={"notes": request.notes}
    )
    
    return ReportResponse.model_validate(report)


@router.post("/{report_id}/unpublish", response_model=ReportResponse)
def unpublish_report(
    report_id: int,
    current_user: User = Depends(require_permission(Permission.PUBLISH_REPORTS.value)),
    db: Session = Depends(get_db)
):
    """Unpublish a report back to DRAFT status."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
    
    from app.models import ReportStatus
    
    if report.status != ReportStatus.PUBLISHED:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Report is not published"
        )
    
    # Revert to draft
    report.status = ReportStatus.DRAFT
    report.published_by_id = None
    report.published_at = None
    
    db.commit()
    db.refresh(report)
    
    logger.info("report_unpublished", report_id=report_id, admin_id=current_user.id)
    
    # Audit log
    AuditManager.log_event(
        db=db,
        user_id=current_user.id,
        event_type=AuditEventType.REPORT_GENERATION,
        action=f"Unpublished report: {report.title}",
        resource_type="report",
        resource_id=report_id
    )
    
    return ReportResponse.model_validate(report)


@router.post("/{report_id}/share")
def share_report(
    report_id: int,
    emails: List[str],
    current_user: User = Depends(require_permission(Permission.SHARE_REPORTS.value)),
    db: Session = Depends(get_db)
):
    """Share a report with email addresses."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
    
    report.shared_with_emails = emails
    db.commit()
    
    logger.info("report_shared", report_id=report_id, recipient_count=len(emails), user_id=current_user.id)
    
    # Audit log
    AuditManager.log_event(
        db=db,
        user_id=current_user.id,
        event_type=AuditEventType.REPORT_GENERATION,
        action=f"Shared report with {len(emails)} recipients",
        resource_type="report",
        resource_id=report_id,
        details={"recipients_count": len(emails), "recipients_sample": emails[:3]}
    )
    
    return {"message": f"Report shared with {len(emails)} recipients"}


async def _generate_executive_summary_genai(articles: List[Article], db: Session) -> str:
    """Generate executive summary using GenAI."""
    from app.genai.provider import get_model_manager
    
    # Collect article summaries and intelligence
    article_data = []
    for article in articles:
        intel = db.query(ExtractedIntelligence).filter(
            ExtractedIntelligence.article_id == article.id
        ).all()
        
        ioc_count = len([i for i in intel if i.intelligence_type.value == "IOC"])
        ttp_count = len([i for i in intel if i.intelligence_type.value == "TTP"])
        
        article_data.append({
            "title": article.title,
            "summary": article.executive_summary or article.summary or "",
            "iocs": ioc_count,
            "ttps": ttp_count,
            "is_high_priority": article.is_high_priority
        })
    
    try:
        manager = get_model_manager()
        result = await manager.generate_with_fallback(
            system_prompt="""You are a threat intelligence analyst writing an executive report for C-level executives.
Create a comprehensive executive summary that:
1. Synthesizes all the threats into key themes
2. Highlights the most critical risks and their business impact
3. Provides strategic recommendations
4. Uses clear, non-technical language

Format the report professionally with clear sections.""",
            user_prompt=f"""Generate an executive summary for these {len(articles)} threat intelligence articles:

{chr(10).join([f"- {a['title']}: {a['summary'][:200]}... (IOCs: {a['iocs']}, TTPs: {a['ttps']}, Priority: {'HIGH' if a['is_high_priority'] else 'Normal'})" for a in article_data])}

Write a cohesive executive summary that synthesizes these threats:"""
        )
        return result.get("response", _generate_executive_summary_fallback(articles))
    except Exception as e:
        logger.warning("genai_executive_summary_failed", error=str(e))
        return _generate_executive_summary_fallback(articles)


def _generate_executive_summary_fallback(articles: List[Article]) -> str:
    """Fallback: Generate executive summary from articles without GenAI."""
    summaries = [a.executive_summary or a.summary for a in articles if a.executive_summary or a.summary]
    return "\n\n".join([f"## {a.title}\n{s}" for a, s in zip(articles, summaries)])


def _generate_executive_summary(articles: List[Article]) -> str:
    """Generate executive summary from articles (sync wrapper)."""
    return _generate_executive_summary_fallback(articles)


async def _generate_technical_summary_genai(articles: List[Article], db: Session) -> str:
    """Generate technical summary using GenAI."""
    from app.genai.provider import get_model_manager
    
    # Collect IOCs and TTPs
    all_iocs = []
    all_ttps = []
    
    for article in articles:
        intel = db.query(ExtractedIntelligence).filter(
            ExtractedIntelligence.article_id == article.id
        ).all()
        
        for i in intel:
            if i.intelligence_type.value == "IOC":
                all_iocs.append({"value": i.value, "type": i.meta.get("type") if i.meta else "unknown", "article": article.title})
            elif i.intelligence_type.value == "TTP":
                all_ttps.append({"mitre_id": i.mitre_id, "name": i.value, "article": article.title})
    
    try:
        manager = get_model_manager()
        result = await manager.generate_with_fallback(
            system_prompt="""You are a senior SOC analyst writing a technical threat report.
Create a detailed technical summary that:
1. Analyzes attack chains and MITRE ATT&CK mappings
2. Lists all IOCs organized by type (IPs, domains, hashes, etc.)
3. Provides detection rules and hunt queries
4. Includes remediation steps

Format professionally with clear sections for SOC analysts.""",
            user_prompt=f"""Generate a technical summary for {len(articles)} articles:

Articles:
{chr(10).join([f"- {a.title}: {(a.technical_summary or a.summary or '')[:200]}..." for a in articles])}

IOCs Found ({len(all_iocs)} total):
{chr(10).join([f"- [{i['type']}] {i['value']}" for i in all_iocs[:30]])}

MITRE ATT&CK TTPs ({len(all_ttps)} total):
{chr(10).join([f"- {t['mitre_id']}: {t['name']}" for t in all_ttps[:20]])}

Write a comprehensive technical summary:"""
        )
        return result.get("response", _generate_technical_summary_fallback(articles))
    except Exception as e:
        logger.warning("genai_technical_summary_failed", error=str(e))
        return _generate_technical_summary_fallback(articles)


def _generate_technical_summary_fallback(articles: List[Article]) -> str:
    """Fallback: Generate technical summary without GenAI."""
    summaries = [a.technical_summary or a.normalized_content for a in articles]
    return "\n\n".join([f"## {a.title}\n{s}" for a, s in zip(articles, summaries) if s])


def _generate_technical_summary(articles: List[Article]) -> str:
    """Generate technical summary from articles (sync wrapper)."""
    return _generate_technical_summary_fallback(articles)


def _extract_key_findings(articles: List[Article]) -> List[str]:
    """Extract key findings from articles."""
    findings = []
    for article in articles:
        if article.is_high_priority:
            findings.append(f"Critical: {article.title}")
        if article.executive_summary:
            # Extract first sentence as a finding
            first_sentence = article.executive_summary.split('.')[0] if '.' in article.executive_summary else article.executive_summary[:100]
            findings.append(first_sentence.strip())
    
    return findings[:10]  # Limit to top 10


def _extract_recommendations(articles: List[Article]) -> List[str]:
    """Extract or generate recommendations from articles."""
    recommendations = [
        "Review and update security controls based on identified threats",
        "Conduct threat hunting for related IOCs across the environment",
        "Update detection rules to identify similar attack patterns",
        "Brief security teams on emerging threat landscape",
        "Schedule follow-up review of remediation actions"
    ]
    
    # Add article-specific recommendations for high priority items
    for article in articles:
        if article.is_high_priority:
            recommendations.insert(0, f"Immediate action required for: {article.title}")
    
    return recommendations[:10]  # Limit to top 10


def _generate_comprehensive_report(articles: List[Article]) -> str:
    """Generate comprehensive report."""
    report = "# Comprehensive Threat Intelligence Report\n\n"
    
    for article in articles:
        report += f"## {article.title}\n"
        report += f"**Status:** {article.status}\n"
        report += f"**Published:** {article.published_at}\n"
        report += f"**Source:** {article.url}\n\n"
        report += f"{article.normalized_content}\n\n"
    
    return report


@router.get("/", response_model=List[ReportResponse])
def list_reports(
    page: int = 1,
    page_size: int = 20,
    current_user: User = Depends(require_permission(Permission.READ_REPORTS.value)),
    db: Session = Depends(get_db)
):
    """List all reports with pagination."""
    offset = (page - 1) * page_size
    reports = db.query(Report).order_by(Report.created_at.desc()).offset(offset).limit(page_size).all()
    return [ReportResponse.model_validate(r) for r in reports]


# ============ DELETE ENDPOINTS FOR REPORTS ============

@router.delete("/{report_id}", status_code=204)
def delete_report(
    report_id: int,
    current_user: User = Depends(require_permission(Permission.CREATE_REPORTS.value)),
    db: Session = Depends(get_db)
):
    """Delete a specific report."""
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
    
    report_title = report.title
    db.delete(report)
    db.commit()
    
    logger.info("report_deleted", report_id=report_id, user_id=current_user.id)
    
    # Audit log
    AuditManager.log_event(
        db=db,
        user_id=current_user.id,
        event_type=AuditEventType.REPORT_GENERATION,
        action=f"Deleted report: {report_title}",
        resource_type="report",
        resource_id=report_id
    )
    
    return None


class BatchDeleteRequest(BaseModel):
    report_ids: List[int]


@router.post("/batch-delete", summary="Delete multiple reports")
def delete_reports_batch(
    request: BatchDeleteRequest,
    current_user: User = Depends(require_permission(Permission.CREATE_REPORTS.value)),
    db: Session = Depends(get_db)
):
    """Delete multiple reports at once."""
    deleted_count = 0
    
    for report_id in request.report_ids:
        report = db.query(Report).filter(Report.id == report_id).first()
        if report:
            db.delete(report)
            deleted_count += 1
    
    db.commit()
    
    logger.info("reports_batch_deleted", count=deleted_count, user_id=current_user.id)
    
    # Audit log
    AuditManager.log_event(
        db=db,
        user_id=current_user.id,
        event_type=AuditEventType.REPORT_GENERATION,
        action=f"Batch deleted {deleted_count} reports",
        resource_type="report",
        details={"deleted_count": deleted_count, "report_ids": request.report_ids}
    )
    
    return {"message": f"Deleted {deleted_count} reports", "deleted_count": deleted_count}


# ============ REVIEWED ARTICLES FOR REPORTS ============

@router.get("/articles/reviewed", summary="Get reviewed articles for report generation")
def get_reviewed_articles_for_reports(
    page: int = Query(1, ge=1),
    page_size: int = Query(50, ge=1, le=200),
    days_back: int = Query(7, ge=1, le=90),
    high_priority_only: bool = Query(False),
    current_user: User = Depends(require_permission(Permission.READ_REPORTS.value)),
    db: Session = Depends(get_db)
):
    """Get reviewed articles ready for report generation.
    
    Returns articles with status REVIEWED, organized by date.
    """
    cutoff_date = datetime.utcnow() - timedelta(days=days_back)
    
    query = db.query(Article).options(
        joinedload(Article.feed_source)
    ).filter(
        Article.status == ArticleStatus.REVIEWED,
        Article.updated_at >= cutoff_date
    )
    
    if high_priority_only:
        query = query.filter(Article.is_high_priority == True)
    
    query = query.order_by(desc(Article.updated_at))
    
    total = query.count()
    articles = query.offset((page - 1) * page_size).limit(page_size).all()
    
    # Group by date
    by_date = {}
    for article in articles:
        date_key = article.updated_at.strftime("%Y-%m-%d")
        if date_key not in by_date:
            by_date[date_key] = []
        
        # Get intelligence count
        intel_count = db.query(ExtractedIntelligence).filter(
            ExtractedIntelligence.article_id == article.id
        ).count()
        
        by_date[date_key].append({
            "id": article.id,
            "title": article.title,
            "source_name": article.feed_source.name if article.feed_source else None,
            "reviewed_at": article.reviewed_at,
            "is_high_priority": article.is_high_priority,
            "watchlist_keywords": article.watchlist_match_keywords or [],
            "intelligence_count": intel_count,
            "executive_summary": article.executive_summary[:200] if article.executive_summary else None,
            "has_analysis": bool(article.executive_summary or article.technical_summary)
        })
    
    return {
        "articles_by_date": by_date,
        "total": total,
        "page": page,
        "page_size": page_size,
        "days_back": days_back
    }


# ============ AUTO-NAMED REPORTS ============

@router.post("/generate/auto", response_model=ReportResponse, summary="Generate auto-named daily/weekly report")
def generate_auto_report(
    request: AutoReportRequest,
    current_user: User = Depends(require_permission(Permission.CREATE_REPORTS.value)),
    db: Session = Depends(get_db)
):
    """Generate an auto-named report from reviewed articles.
    
    - Daily reports: Include articles reviewed today
    - Weekly reports: Include articles reviewed in the last 7 days
    
    Report is auto-named as: YYYY-MM-DD_type (e.g., 2026-01-16_executive)
    """
    today = datetime.utcnow().date()
    
    if request.period == "daily":
        start_date = datetime.combine(today, datetime.min.time())
        end_date = datetime.combine(today, datetime.max.time())
        period_str = today.strftime("%Y-%m-%d")
    else:  # weekly
        start_date = datetime.combine(today - timedelta(days=7), datetime.min.time())
        end_date = datetime.combine(today, datetime.max.time())
        week_start = (today - timedelta(days=7)).strftime("%Y-%m-%d")
        period_str = f"{week_start}_to_{today.strftime('%Y-%m-%d')}"
    
    # Query reviewed articles in the period
    query = db.query(Article).filter(
        Article.status == ArticleStatus.REVIEWED,
        Article.updated_at >= start_date,
        Article.updated_at <= end_date
    )
    
    if request.include_high_priority_only:
        query = query.filter(Article.is_high_priority == True)
    
    articles = query.order_by(desc(Article.updated_at)).all()
    
    if not articles:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"No reviewed articles found for {request.period} period"
        )
    
    # Auto-generate title
    title = f"{period_str}_{request.report_type}"
    if request.include_high_priority_only:
        title += "_high_priority"
    
    # Generate content
    if request.report_type == "executive":
        content = _generate_executive_summary(articles)
    elif request.report_type == "technical":
        content = _generate_technical_summary(articles)
    else:
        content = _generate_comprehensive_report(articles)
    
    # Create report
    report = Report(
        title=title,
        article_ids=[a.id for a in articles],
        content=content,
        report_type=request.report_type,
        generated_by_id=current_user.id,
        generated_at=datetime.utcnow()
    )
    
    db.add(report)
    db.commit()
    db.refresh(report)
    
    logger.info(
        "auto_report_generated",
        report_id=report.id,
        title=title,
        period=request.period,
        article_count=len(articles),
        user_id=current_user.id
    )
    
    # Audit log
    AuditManager.log_event(
        db=db,
        user_id=current_user.id,
        event_type=AuditEventType.REPORT_GENERATION,
        action=f"Auto-generated {request.period} {request.report_type} report: {title}",
        resource_type="report",
        resource_id=report.id,
        details={"period": request.period, "articles_count": len(articles), "high_priority_only": request.include_high_priority_only}
    )
    
    return ReportResponse.model_validate(report)


@router.get("/stats", summary="Get report statistics")
def get_report_stats(
    current_user: User = Depends(require_permission(Permission.READ_REPORTS.value)),
    db: Session = Depends(get_db)
):
    """Get statistics about reports and reviewed articles."""
    today = datetime.utcnow().date()
    week_ago = today - timedelta(days=7)
    
    total_reports = db.query(Report).count()
    
    # Reviewed articles counts
    total_reviewed = db.query(Article).filter(
        Article.status == ArticleStatus.REVIEWED
    ).count()
    
    reviewed_today = db.query(Article).filter(
        Article.status == ArticleStatus.REVIEWED,
        Article.updated_at >= datetime.combine(today, datetime.min.time())
    ).count()
    
    reviewed_this_week = db.query(Article).filter(
        Article.status == ArticleStatus.REVIEWED,
        Article.updated_at >= datetime.combine(week_ago, datetime.min.time())
    ).count()
    
    high_priority_reviewed = db.query(Article).filter(
        Article.status == ArticleStatus.REVIEWED,
        Article.is_high_priority == True
    ).count()
    
    # Recent reports
    recent_reports = db.query(Report).order_by(desc(Report.created_at)).limit(5).all()
    
    return {
        "total_reports": total_reports,
        "reviewed_articles": {
            "total": total_reviewed,
            "today": reviewed_today,
            "this_week": reviewed_this_week,
            "high_priority": high_priority_reviewed
        },
        "recent_reports": [
            {
                "id": r.id,
                "title": r.title,
                "type": r.report_type,
                "article_count": len(r.article_ids),
                "generated_at": r.generated_at
            }
            for r in recent_reports
        ]
    }


@router.get("/{report_id}/export/csv")
def export_report_csv(
    report_id: int,
    current_user: User = Depends(require_permission(Permission.READ_REPORTS.value)),
    db: Session = Depends(get_db)
):
    """Export report data as CSV.
    
    Includes all articles, IOCs, TTPs, hunt results, and GenAI analysis.
    """
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")

    def _csv_safe(value):
        """
        Prevent CSV/Excel formula injection by prefixing dangerous leading characters.
        See: OWASP CSV Injection (formula injection).
        """
        if value is None:
            return ""
        if isinstance(value, (int, float, bool)):
            return value
        s = str(value)
        if not s:
            return s
        if s.lstrip().startswith(("=", "+", "-", "@")):
            return "'" + s
        return s
    
    # Fetch articles
    articles = db.query(Article).filter(Article.id.in_(report.article_ids)).all()
    
    # Create CSV in memory
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Write header
    writer.writerow([
        "Report Title", "Report Type", "Generated At",
        "Article ID", "Article Title", "Article Status", "Published At", "Source URL",
        "Intelligence Type", "Value", "Confidence", "MITRE ID",
        "Hunt Platform", "Hunt Status", "Results Count", "Execution Time (ms)",
        "Risk Level", "AI Analysis Summary"
    ])
    
    # Write data
    for article in articles:
        # Get intelligence for this article
        intelligence = db.query(ExtractedIntelligence).filter(
            ExtractedIntelligence.article_id == article.id
        ).all()
        
        # Get hunts for this article
        hunts = db.query(Hunt).filter(Hunt.article_id == article.id).all()
        
        # If no intelligence, still write article info
        if not intelligence and not hunts:
            writer.writerow([
                _csv_safe(report.title), _csv_safe(report.report_type), report.generated_at.isoformat(),
                article.id, _csv_safe(article.title), _csv_safe(article.status.value) if article.status else "",
                article.published_at.isoformat() if article.published_at else "", _csv_safe(article.url),
                "", "", "", "",
                "", "", "", "",
                "", ""
            ])
        
        # Write intelligence rows
        for intel in intelligence:
            writer.writerow([
                _csv_safe(report.title), _csv_safe(report.report_type), report.generated_at.isoformat(),
                article.id, _csv_safe(article.title), _csv_safe(article.status.value) if article.status else "",
                article.published_at.isoformat() if article.published_at else "", _csv_safe(article.url),
                _csv_safe(intel.intelligence_type.value) if intel.intelligence_type else "",
                _csv_safe(intel.value), intel.confidence, _csv_safe(intel.mitre_id or ""),
                "", "", "", "",
                "", ""
            ])
        
        # Write hunt execution rows
        for hunt in hunts:
            for execution in hunt.executions:
                analysis = execution.results.get("genai_analysis", {}) if execution.results else {}
                writer.writerow([
                    _csv_safe(report.title), _csv_safe(report.report_type), report.generated_at.isoformat(),
                    article.id, _csv_safe(article.title), _csv_safe(article.status.value) if article.status else "",
                    article.published_at.isoformat() if article.published_at else "", _csv_safe(article.url),
                    "", "", "", "",
                    _csv_safe(hunt.platform), _csv_safe(execution.status.value) if execution.status else "",
                    execution.results.get("results_count", 0) if execution.results else 0,
                    execution.execution_time_ms or 0,
                    _csv_safe(analysis.get("risk_level", "")),
                    _csv_safe((analysis.get("executive_summary", "") or "")[:200])
                ])
    
    output.seek(0)
    
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=report_{report_id}.csv"}
    )


@router.get("/{report_id}/export/docx")
def export_report_docx(
    report_id: int,
    current_user: User = Depends(require_permission(Permission.READ_REPORTS.value)),
    db: Session = Depends(get_db)
):
    """Export report as Word document (.docx).
    
    Creates a professionally formatted threat intelligence report.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_501_NOT_IMPLEMENTED,
            detail="python-docx not installed. Run: pip install python-docx"
        )
    
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
    
    # Fetch articles
    articles = db.query(Article).filter(Article.id.in_(report.article_ids)).all()
    
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading(report.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f"Report Type: {report.report_type.title()}")
    doc.add_paragraph(f"Generated: {report.generated_at.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Articles Included: {len(articles)}")
    doc.add_paragraph()
    
    # Table of Contents
    doc.add_heading("Table of Contents", level=1)
    for i, article in enumerate(articles, 1):
        doc.add_paragraph(f"{i}. {article.title[:80]}", style="List Number")
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    exec_summary_parts = []
    for article in articles:
        if article.executive_summary:
            exec_summary_parts.append(f"• {article.title}: {article.executive_summary[:200]}...")
    
    if exec_summary_parts:
        doc.add_paragraph("\n".join(exec_summary_parts))
    else:
        doc.add_paragraph("No executive summaries available for the selected articles.")
    doc.add_paragraph()
    
    # Detailed Analysis
    doc.add_heading("Detailed Analysis", level=1)
    
    for article in articles:
        # Article section
        doc.add_heading(article.title, level=2)
        
        # Metadata table
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        
        cells = table.rows[0].cells
        cells[0].text = "Status"
        cells[1].text = article.status.value if article.status else "N/A"
        
        cells = table.rows[1].cells
        cells[0].text = "Published"
        cells[1].text = article.published_at.strftime("%Y-%m-%d") if article.published_at else "N/A"
        
        cells = table.rows[2].cells
        cells[0].text = "Source"
        cells[1].text = article.url or "N/A"
        
        cells = table.rows[3].cells
        cells[0].text = "Priority"
        cells[1].text = "High" if article.is_high_priority else "Normal"
        
        doc.add_paragraph()
        
        # Summary
        if article.summary:
            doc.add_heading("Summary", level=3)
            doc.add_paragraph(article.summary)
        
        # Extracted Intelligence
        intelligence = db.query(ExtractedIntelligence).filter(
            ExtractedIntelligence.article_id == article.id
        ).all()
        
        if intelligence:
            doc.add_heading("Extracted Intelligence", level=3)
            
            # Group by type
            iocs = [i for i in intelligence if i.intelligence_type.value == "IOC"]
            ttps = [i for i in intelligence if i.intelligence_type.value == "TTP"]
            ioas = [i for i in intelligence if i.intelligence_type.value == "IOA"]
            
            if iocs:
                doc.add_paragraph("IOCs (Indicators of Compromise):", style="Intense Quote")
                for ioc in iocs[:10]:  # Limit to 10
                    doc.add_paragraph(f"• {ioc.value} (Confidence: {ioc.confidence}%)", style="List Bullet")
            
            if ttps:
                doc.add_paragraph("TTPs (MITRE ATT&CK):", style="Intense Quote")
                for ttp in ttps:
                    doc.add_paragraph(f"• {ttp.mitre_id}: {ttp.value}", style="List Bullet")
            
            if ioas:
                doc.add_paragraph("IOAs (Indicators of Attack):", style="Intense Quote")
                for ioa in ioas[:5]:
                    doc.add_paragraph(f"• {ioa.value}", style="List Bullet")
        
        # Hunt Results
        hunts = db.query(Hunt).filter(Hunt.article_id == article.id).all()
        
        if hunts:
            doc.add_heading("Hunt Results", level=3)
            
            for hunt in hunts:
                doc.add_paragraph(f"Platform: {hunt.platform.upper()}")
                
                for execution in hunt.executions:
                    status_str = execution.status.value if execution.status else "Unknown"
                    results_count = execution.results.get("results_count", 0) if execution.results else 0
                    
                    doc.add_paragraph(
                        f"  Status: {status_str} | Hits: {results_count} | "
                        f"Time: {execution.execution_time_ms or 0}ms"
                    )
                    
                    # GenAI Analysis
                    if execution.results and execution.results.get("genai_analysis"):
                        analysis = execution.results["genai_analysis"]
                        doc.add_paragraph("AI Analysis:", style="Intense Quote")
                        doc.add_paragraph(f"Risk Level: {analysis.get('risk_level', 'N/A')}")
                        if analysis.get("executive_summary"):
                            doc.add_paragraph(analysis["executive_summary"])
                        if analysis.get("recommended_actions"):
                            doc.add_paragraph("Recommended Actions:")
                            for action in analysis["recommended_actions"][:3]:
                                doc.add_paragraph(f"• {action}", style="List Bullet")
        
        doc.add_page_break()
    
    # Appendix - IOC List
    doc.add_heading("Appendix A: Complete IOC List", level=1)
    
    all_iocs = db.query(ExtractedIntelligence).filter(
        ExtractedIntelligence.article_id.in_(report.article_ids),
        ExtractedIntelligence.intelligence_type == "IOC"
    ).all()
    
    if all_iocs:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Type"
        hdr_cells[1].text = "Value"
        hdr_cells[2].text = "Confidence"
        hdr_cells[3].text = "Source Article"
        
        for ioc in all_iocs[:100]:  # Limit to 100
            row_cells = table.add_row().cells
            row_cells[0].text = str(ioc.meta.get("type", "unknown") if ioc.meta else "unknown")
            row_cells[1].text = ioc.value[:50]
            row_cells[2].text = f"{ioc.confidence}%"
            row_cells[3].text = str(ioc.article_id)
    else:
        doc.add_paragraph("No IOCs extracted from the selected articles.")
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return StreamingResponse(
        doc_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=report_{report_id}.docx"}
    )


@router.get("/{report_id}/export/pdf")
def export_report_pdf(
    report_id: int,
    current_user: User = Depends(require_permission(Permission.READ_REPORTS.value)),
    db: Session = Depends(get_db)
):
    """Export report as professionally styled PDF document.
    
    Creates a premium formatted threat intelligence report in PDF format.
    Matches Analytics Dashboard styling with company branding support.
    Uses ReportLab for PDF generation (pure Python, no system deps).
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, ListFlowable, ListItem, Image, Flowable
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT
    from reportlab.graphics.shapes import Drawing, Rect
    from reportlab.graphics import renderPDF
    import urllib.request
    
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
    
    # Fetch articles
    articles = db.query(Article).filter(Article.id.in_(report.article_ids)).all()
    
    # Get branding settings from SystemConfiguration
    branding = {}
    try:
        from app.models import SystemConfiguration
        configs = db.query(SystemConfiguration).filter(
            SystemConfiguration.category == 'branding'
        ).all()
        for config in configs:
            branding[config.key] = config.value
    except:
        pass
    
    company_name = branding.get('company_name', 'Parshu Security')
    company_logo_url = branding.get('company_logo_url', '')
    confidentiality = branding.get('confidentiality_notice', 'CONFIDENTIAL - Internal Use Only')
    report_footer = branding.get('report_footer', 'Generated by Parshu Threat Intelligence Platform')
    
    # Create PDF in memory
    pdf_buffer = io.BytesIO()
    
    # Refined professional color palette
    PRIMARY_DARK = '#1e3a5f'      # Deep navy blue (more professional than dark purple)
    SECONDARY_DARK = '#2c5282'    # Medium navy
    PRIMARY_BLUE = '#3182ce'      # Professional blue (softer than bright blue)
    ACCENT_BLUE = '#4299e1'       # Light accent blue
    SUCCESS_GREEN = '#38a169'     # Muted green (more professional)
    DANGER_RED = '#c53030'        # Muted red (less aggressive)
    WARNING_ORANGE = '#dd6b20'    # Muted orange
    TEXT_DARK = '#1a202c'         # Near black for main text
    TEXT_BODY = '#2d3748'         # Dark gray for body text
    TEXT_LIGHT = '#4a5568'        # Medium gray for secondary text
    TEXT_MUTED = '#718096'        # Light gray for muted text
    BG_LIGHT = '#f7fafc'          # Very light gray background
    BG_SECTION = '#edf2f7'        # Slightly darker section background
    BORDER_LIGHT = '#e2e8f0'      # Light border
    BORDER_MEDIUM = '#cbd5e0'     # Medium border
    
    # Custom flowable for colored header bar
    class HeaderBar(Flowable):
        def __init__(self, width, height, color):
            Flowable.__init__(self)
            self.width = width
            self.height = height
            self.color = color
            
        def draw(self):
            self.canv.setFillColor(colors.HexColor(self.color))
            self.canv.rect(0, 0, self.width, self.height, fill=1, stroke=0)
    
    # Custom flowable for gradient simulation (solid dark)
    class DarkHeader(Flowable):
        def __init__(self, width, height, title, subtitle, logo_data=None):
            Flowable.__init__(self)
            self.width = width
            self.height = height
            self.title = title
            self.subtitle = subtitle
            self.logo_data = logo_data
            
        def draw(self):
            # Dark background
            self.canv.setFillColor(colors.HexColor(PRIMARY_DARK))
            self.canv.rect(0, 0, self.width, self.height, fill=1, stroke=0)
            
            # Company name
            self.canv.setFillColor(colors.white)
            self.canv.setFont('Helvetica-Bold', 24)
            self.canv.drawString(20, self.height - 40, company_name)
            
            # Title
            self.canv.setFont('Helvetica', 14)
            self.canv.drawString(20, self.height - 60, self.title)
            
            # Subtitle/date on right
            self.canv.setFont('Helvetica', 10)
            self.canv.drawRightString(self.width - 20, self.height - 40, self.subtitle)
    
    # Custom page template with professional header/footer
    def add_page_header_footer(canvas, doc):
        canvas.saveState()
        
        # Top header bar (dark)
        canvas.setFillColor(colors.HexColor(PRIMARY_DARK))
        canvas.rect(0, 730, 612, 62, fill=1, stroke=0)
        
        # Company name in header
        canvas.setFillColor(colors.white)
        canvas.setFont('Helvetica-Bold', 12)
        canvas.drawString(72, 755, company_name.upper())
        
        # Report title in header
        canvas.setFont('Helvetica', 9)
        canvas.drawString(72, 740, f"Threat Intelligence Report")
        
        # Page number on right
        canvas.drawRightString(540, 748, f"Page {doc.page}")
        
        # Blue accent line
        canvas.setStrokeColor(colors.HexColor(PRIMARY_BLUE))
        canvas.setLineWidth(3)
        canvas.line(0, 730, 612, 730)
        
        # Footer
        canvas.setFillColor(colors.HexColor(PRIMARY_DARK))
        canvas.rect(0, 0, 612, 40, fill=1, stroke=0)
        
        canvas.setFillColor(colors.white)
        canvas.setFont('Helvetica', 8)
        canvas.drawString(72, 18, f"Generated: {report.generated_at.strftime('%B %d, %Y at %H:%M UTC')}")
        canvas.drawCentredString(306, 18, confidentiality)
        canvas.drawRightString(540, 18, "PARSHU PLATFORM")
        
        canvas.restoreState()
    
    doc = SimpleDocTemplate(
        pdf_buffer, 
        pagesize=letter, 
        rightMargin=50, 
        leftMargin=50, 
        topMargin=100, 
        bottomMargin=60
    )
    
    # Styles
    styles = getSampleStyleSheet()
    
    # Professional title style (large, centered)
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=28,
        textColor=colors.HexColor(PRIMARY_DARK),
        spaceAfter=8,
        spaceBefore=20,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold',
        leading=32
    )
    
    subtitle_style = ParagraphStyle(
        'Subtitle',
        parent=styles['Normal'],
        fontSize=14,
        textColor=colors.HexColor(TEXT_LIGHT),
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica'
    )
    
    # Section header with blue left border simulation
    section_header_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontSize=15,
        textColor=colors.HexColor(PRIMARY_DARK),
        spaceBefore=20,
        spaceAfter=14,
        fontName='Helvetica-Bold',
        borderPadding=8,
        leftIndent=0
    )
    
    heading2_style = ParagraphStyle(
        'CustomH2',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor(PRIMARY_DARK),
        spaceBefore=18,
        spaceAfter=10,
        fontName='Helvetica-Bold'
    )
    
    heading3_style = ParagraphStyle(
        'CustomH3',
        parent=styles['Heading3'],
        fontSize=12,
        textColor=colors.HexColor(PRIMARY_BLUE),
        spaceBefore=12,
        spaceAfter=6,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor(TEXT_BODY),
        spaceAfter=10,
        spaceBefore=4,
        alignment=TA_JUSTIFY,
        leading=16,
        leftIndent=0,
        firstLineIndent=0
    )
    
    # Indented body for content within sections
    body_indented_style = ParagraphStyle(
        'BodyIndented',
        parent=body_style,
        leftIndent=15,
        rightIndent=10
    )
    
    high_priority_style = ParagraphStyle(
        'HighPriority',
        parent=heading3_style,
        textColor=colors.HexColor(DANGER_RED),
        fontName='Helvetica-Bold'
    )
    
    metric_value_style = ParagraphStyle(
        'MetricValue',
        parent=styles['Normal'],
        fontSize=24,
        textColor=colors.HexColor(PRIMARY_BLUE),
        alignment=TA_CENTER,
        fontName='Helvetica-Bold',
        spaceAfter=4
    )
    
    metric_label_style = ParagraphStyle(
        'MetricLabel',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.HexColor(TEXT_LIGHT),
        alignment=TA_CENTER,
        fontName='Helvetica',
        textTransform='uppercase'
    )
    
    # Build content
    story = []
    
    # ========== COVER PAGE ==========
    # Confidentiality Banner
    conf_banner = Table(
        [[confidentiality.upper()]],
        colWidths=[6*inch]
    )
    conf_banner.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(DANGER_RED)),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.white),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(conf_banner)
    story.append(Spacer(1, 40))
    
    # Company Name (large)
    story.append(Paragraph(company_name.upper(), ParagraphStyle(
        'CompanyName',
        parent=styles['Normal'],
        fontSize=14,
        textColor=colors.HexColor(TEXT_LIGHT),
        alignment=TA_CENTER,
        fontName='Helvetica',
        letterSpacing=4
    )))
    story.append(Spacer(1, 10))
    
    # Report Title
    story.append(Paragraph(report.title, title_style))
    story.append(Paragraph("Threat Intelligence Report", subtitle_style))
    story.append(Spacer(1, 20))
    
    # Metrics Summary Cards (4 columns)
    high_priority_count = len([a for a in articles if a.is_high_priority])
    total_intel = sum([
        db.query(ExtractedIntelligence).filter(ExtractedIntelligence.article_id == a.id).count()
        for a in articles
    ])
    ioc_count = sum([
        db.query(ExtractedIntelligence).filter(
            ExtractedIntelligence.article_id == a.id,
            ExtractedIntelligence.intelligence_type == 'IOC'
        ).count()
        for a in articles
    ])
    ttp_count = sum([
        db.query(ExtractedIntelligence).filter(
            ExtractedIntelligence.article_id == a.id,
            ExtractedIntelligence.intelligence_type == 'TTP'
        ).count()
        for a in articles
    ])
    
    # Metrics table styled as cards
    metrics_data = [
        [str(len(articles)), str(high_priority_count), str(ioc_count), str(ttp_count)],
        ['ARTICLES', 'HIGH PRIORITY', 'IOCs', 'TTPs']
    ]
    
    metrics_table = Table(metrics_data, colWidths=[1.25*inch, 1.25*inch, 1.25*inch, 1.25*inch])
    metrics_table.setStyle(TableStyle([
        # First row - values (refined colors)
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 26),
        ('TEXTCOLOR', (0, 0), (0, 0), colors.HexColor(PRIMARY_BLUE)),
        ('TEXTCOLOR', (1, 0), (1, 0), colors.HexColor(DANGER_RED)),
        ('TEXTCOLOR', (2, 0), (2, 0), colors.HexColor(SUCCESS_GREEN)),
        ('TEXTCOLOR', (3, 0), (3, 0), colors.HexColor(WARNING_ORANGE)),
        # Second row - labels
        ('FONTNAME', (0, 1), (-1, 1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, 1), 8),
        ('TEXTCOLOR', (0, 1), (-1, 1), colors.HexColor(TEXT_MUTED)),
        # All cells
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BACKGROUND', (0, 0), (-1, -1), colors.white),
        ('TOPPADDING', (0, 0), (-1, 0), 18),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
        ('TOPPADDING', (0, 1), (-1, 1), 4),
        ('BOTTOMPADDING', (0, 1), (-1, 1), 18),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor(BORDER_MEDIUM)),
        ('LINEBEFORE', (1, 0), (1, -1), 0.5, colors.HexColor(BORDER_LIGHT)),
        ('LINEBEFORE', (2, 0), (2, -1), 0.5, colors.HexColor(BORDER_LIGHT)),
        ('LINEBEFORE', (3, 0), (3, -1), 0.5, colors.HexColor(BORDER_LIGHT)),
    ]))
    story.append(metrics_table)
    story.append(Spacer(1, 30))
    
    # Report Metadata
    meta_data = [
        ['Report Type', report.report_type.replace('_', ' ').title()],
        ['Generated', report.generated_at.strftime('%B %d, %Y at %H:%M UTC')],
        ['Articles Analyzed', str(len(articles))],
        ['Classification', confidentiality.split('-')[0].strip()],
        ['Prepared By', current_user.username]
    ]
    
    meta_table = Table(meta_data, colWidths=[1.8*inch, 3.5*inch])
    meta_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor(PRIMARY_DARK)),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.white),
        ('TEXTCOLOR', (1, 0), (1, -1), colors.HexColor(TEXT_DARK)),
        ('BACKGROUND', (1, 0), (1, -1), colors.white),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 10),
        ('LEFTPADDING', (0, 0), (-1, -1), 12),
        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor(PRIMARY_DARK)),
        ('LINEBELOW', (0, 0), (-1, -2), 0.5, colors.HexColor(BORDER_LIGHT)),
    ]))
    story.append(meta_table)
    story.append(PageBreak())
    
    # ========== EXECUTIVE SUMMARY SECTION ==========
    # Section header with refined styling
    exec_header = Table(
        [['', 'Executive Summary']],
        colWidths=[4, 5.2*inch]
    )
    exec_header.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, 0), colors.HexColor(PRIMARY_BLUE)),
        ('BACKGROUND', (1, 0), (1, 0), colors.HexColor(BG_SECTION)),
        ('TEXTCOLOR', (1, 0), (1, 0), colors.HexColor(PRIMARY_DARK)),
        ('FONTNAME', (1, 0), (1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (1, 0), (1, 0), 13),
        ('LEFTPADDING', (1, 0), (1, 0), 14),
        ('TOPPADDING', (0, 0), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(exec_header)
    story.append(Spacer(1, 14))
    
    # Summary content
    exec_summaries = []
    for article in articles:
        summary = article.executive_summary or article.summary
        if summary:
            priority_tag = '<font color="#ff4d4f">[HIGH PRIORITY]</font> ' if article.is_high_priority else ''
            exec_summaries.append(f"{priority_tag}<b>{article.title}:</b> {summary[:350]}...")
    
    if exec_summaries:
        for idx, summary in enumerate(exec_summaries):
            # Create a card-like summary with left accent border
            summary_content = Paragraph(summary, body_indented_style)
            summary_box = Table(
                [['', summary_content]],
                colWidths=[4, 5*inch]
            )
            summary_box.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, 0), colors.HexColor(ACCENT_BLUE)),  # Left accent bar
                ('BACKGROUND', (1, 0), (1, 0), colors.white),
                ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor(BORDER_MEDIUM)),
                ('LEFTPADDING', (1, 0), (1, 0), 15),
                ('RIGHTPADDING', (1, 0), (1, 0), 15),
                ('TOPPADDING', (0, 0), (-1, -1), 12),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(summary_box)
            story.append(Spacer(1, 10))
    else:
        story.append(Paragraph("No executive summaries available for the selected articles.", body_style))
    
    story.append(PageBreak())
    
    # ========== DETAILED ANALYSIS SECTION ==========
    detail_header = Table(
        [['', 'Detailed Analysis']],
        colWidths=[4, 5.2*inch]
    )
    detail_header.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, 0), colors.HexColor(SUCCESS_GREEN)),
        ('BACKGROUND', (1, 0), (1, 0), colors.HexColor(BG_SECTION)),
        ('TEXTCOLOR', (1, 0), (1, 0), colors.HexColor(PRIMARY_DARK)),
        ('FONTNAME', (1, 0), (1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (1, 0), (1, 0), 13),
        ('LEFTPADDING', (1, 0), (1, 0), 14),
        ('TOPPADDING', (0, 0), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(detail_header)
    story.append(Spacer(1, 16))
    
    for idx, article in enumerate(articles):
        # Article card header with refined styling
        priority_color = DANGER_RED if article.is_high_priority else SECONDARY_DARK
        priority_text = 'HIGH PRIORITY' if article.is_high_priority else f'Article {idx + 1}'
        
        article_header = Table(
            [[priority_text, article.title[:55] + ('...' if len(article.title) > 55 else '')]],
            colWidths=[1.1*inch, 4.1*inch]
        )
        article_header.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, 0), colors.HexColor(priority_color)),
            ('TEXTCOLOR', (0, 0), (0, 0), colors.white),
            ('BACKGROUND', (1, 0), (1, 0), colors.white),
            ('TEXTCOLOR', (1, 0), (1, 0), colors.HexColor(TEXT_DARK)),
            ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (0, 0), 8),
            ('FONTSIZE', (1, 0), (1, 0), 10),
            ('LEFTPADDING', (0, 0), (-1, -1), 12),
            ('RIGHTPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor(BORDER_MEDIUM)),
        ]))
        story.append(article_header)
        
        # Article metadata
        article_meta = [
            ['Status', article.status.value.replace('_', ' ').title() if article.status else 'N/A',
             'Published', article.published_at.strftime('%Y-%m-%d') if article.published_at else 'N/A']
        ]
        
        meta_table = Table(article_meta, colWidths=[0.8*inch, 1.8*inch, 0.8*inch, 1.8*inch])
        meta_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor(TEXT_LIGHT)),
            ('BACKGROUND', (0, 0), (-1, -1), colors.white),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ]))
        story.append(meta_table)
        
        # Summary with proper indentation
        if article.summary:
            story.append(Spacer(1, 8))
            summary_para = Paragraph(article.summary[:600], body_indented_style)
            story.append(summary_para)
        
        # Intelligence
        intelligence = db.query(ExtractedIntelligence).filter(
            ExtractedIntelligence.article_id == article.id
        ).all()
        
        if intelligence:
            iocs = [i for i in intelligence if i.intelligence_type.value == "IOC"]
            ttps = [i for i in intelligence if i.intelligence_type.value == "TTP"]
            
            if iocs:
                story.append(Spacer(1, 12))
                ioc_header = Paragraph(f'<font color="{PRIMARY_BLUE}"><b>Indicators of Compromise ({len(iocs[:10])})</b></font>', body_indented_style)
                story.append(ioc_header)
                story.append(Spacer(1, 4))
                
                ioc_data = [['Type', 'Value', 'Conf.']]
                for ioc in iocs[:10]:
                    ioc_type = ioc.meta.get("type", "unknown").upper() if ioc.meta else "UNKNOWN"
                    ioc_data.append([ioc_type, ioc.value[:42], f"{ioc.confidence}%"])
                
                ioc_table = Table(ioc_data, colWidths=[0.9*inch, 3.7*inch, 0.6*inch])
                ioc_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(SECONDARY_DARK)),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 8),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                    ('FONTNAME', (0, 1), (-1, -1), 'Courier'),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor(TEXT_BODY)),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor(BG_LIGHT)]),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor(BORDER_LIGHT)),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('LEFTPADDING', (0, 0), (-1, -1), 10),
                    ('ALIGN', (-1, 0), (-1, -1), 'CENTER'),
                ]))
                story.append(ioc_table)
            
            if ttps:
                story.append(Spacer(1, 12))
                ttp_header = Paragraph(f'<font color="{WARNING_ORANGE}"><b>MITRE ATT&CK TTPs ({len(ttps[:8])})</b></font>', body_indented_style)
                story.append(ttp_header)
                story.append(Spacer(1, 4))
                
                ttp_data = [['Technique ID', 'Technique Name']]
                for ttp in ttps[:8]:
                    ttp_data.append([ttp.mitre_id or 'N/A', ttp.value[:48]])
                
                ttp_table = Table(ttp_data, colWidths=[1.2*inch, 4*inch])
                ttp_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(WARNING_ORANGE)),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 8),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor(TEXT_BODY)),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor(BG_LIGHT)]),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor(BORDER_LIGHT)),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('LEFTPADDING', (0, 0), (-1, -1), 10),
                ]))
                story.append(ttp_table)
        
        story.append(Spacer(1, 20))
    
    # ========== SOURCES & REFERENCES SECTION ==========
    story.append(PageBreak())
    
    sources_header = Table(
        [['', 'Sources & References']],
        colWidths=[4, 5.2*inch]
    )
    sources_header.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, 0), colors.HexColor(WARNING_ORANGE)),
        ('BACKGROUND', (1, 0), (1, 0), colors.HexColor(BG_SECTION)),
        ('TEXTCOLOR', (1, 0), (1, 0), colors.HexColor(PRIMARY_DARK)),
        ('FONTNAME', (1, 0), (1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (1, 0), (1, 0), 13),
        ('LEFTPADDING', (1, 0), (1, 0), 14),
        ('TOPPADDING', (0, 0), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(sources_header)
    story.append(Spacer(1, 16))
    
    sources_data = [['#', 'ARTICLE TITLE', 'SOURCE URL']]
    for i, article in enumerate(articles, 1):
        sources_data.append([
            str(i), 
            article.title[:55] + ('...' if len(article.title) > 55 else ''),
            (article.url or 'N/A')[:45] + ('...' if article.url and len(article.url) > 45 else '')
        ])
    
    sources_table = Table(sources_data, colWidths=[0.4*inch, 3*inch, 2*inch])
    sources_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(PRIMARY_DARK)),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor(TEXT_DARK)),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor(BORDER_LIGHT)),
        ('ALIGN', (0, 0), (0, -1), 'CENTER'),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor(BG_LIGHT)]),
    ]))
    story.append(sources_table)
    
    # ========== DISCLAIMER ==========
    story.append(Spacer(1, 40))
    
    disclaimer_box = Table(
        [[Paragraph(
            f"<b>DISCLAIMER:</b> This report was generated by {company_name} using the Parshu Threat Intelligence Platform. "
            f"The information contained herein is classified as {confidentiality.split('-')[0].strip()} and is intended for authorized recipients only. "
            "Unauthorized distribution or disclosure is strictly prohibited. The intelligence provided should be validated "
            "before operational use.",
            ParagraphStyle('Disclaimer', parent=body_style, fontSize=8, textColor=colors.HexColor(TEXT_MUTED), alignment=TA_CENTER)
        )]],
        colWidths=[5.4*inch]
    )
    disclaimer_box.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(BG_LIGHT)),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor(BORDER_LIGHT)),
        ('TOPPADDING', (0, 0), (-1, -1), 15),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 15),
        ('LEFTPADDING', (0, 0), (-1, -1), 15),
        ('RIGHTPADDING', (0, 0), (-1, -1), 15),
    ]))
    story.append(disclaimer_box)
    
    # Build PDF with header/footer
    doc.build(story, onFirstPage=add_page_header_footer, onLaterPages=add_page_header_footer)
    pdf_buffer.seek(0)
    
    # Generate professional filename
    safe_title = report.title.replace(' ', '_').replace('/', '-')[:30]
    date_str = report.generated_at.strftime('%Y%m%d')
    filename = f"Parshu_Report_{safe_title}_{date_str}.pdf"
    
    logger.info("pdf_report_exported", report_id=report_id, user_id=current_user.id)
    
    return StreamingResponse(
        pdf_buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


def _generate_pdf_executive_summary(articles: List[Article], db: Session) -> str:
    """Generate executive summary for PDF."""
    summaries = []
    for article in articles:
        summary = article.executive_summary or article.summary or "No summary available."
        summaries.append(f"<strong>{article.title}:</strong> {summary[:300]}...")
    
    return "<br><br>".join(summaries) if summaries else "No executive summaries available."


@router.get("/{report_id}/export/html")
def export_report_html(
    report_id: int,
    current_user: User = Depends(require_permission(Permission.READ_REPORTS.value)),
    db: Session = Depends(get_db)
):
    """Export report as HTML document for browser viewing.
    
    Creates a professionally formatted, browser-viewable threat intelligence report.
    Compatible with Word, PDF printing, and direct viewing.
    """
    import re
    import html
    from urllib.parse import urlparse
    
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
    
    # Fetch articles for additional context
    articles = db.query(Article).filter(Article.id.in_(report.article_ids)).all()
    
    # Get branding settings
    branding = {}
    try:
        from app.models import SystemConfiguration
        configs = db.query(SystemConfiguration).filter(
            SystemConfiguration.category == 'branding'
        ).all()
        for config in configs:
            branding[config.key] = config.value
    except:
        pass
    
    company_name = branding.get('company_name', 'Parshu Security')
    company_logo_url = branding.get('company_logo_url', '')
    confidentiality = branding.get('confidentiality_notice', 'CONFIDENTIAL - Internal Use Only')
    primary_color = branding.get('primary_color', '#1890ff')
    dark_color = branding.get('dark_color', '#1a1a2e')

    def _sanitize_css_color(value: str, fallback: str) -> str:
        v = (value or "").strip()
        if re.match(r"^#[0-9a-fA-F]{3,8}$", v):
            return v
        if re.match(r"^(rgb|rgba|hsl|hsla)\\([0-9\\s,%.]+\\)$", v):
            return v
        return fallback

    def _sanitize_href(url: str) -> str:
        if not url:
            return "#"
        u = str(url).strip()
        # Reject obviously unsafe characters that can lead to HTML/attribute injection
        if re.search(r"[\x00-\x1f\x7f\s\"'<>]", u):
            return "#"
        try:
            parsed = urlparse(u)
        except Exception:
            return "#"
        if parsed.scheme in ("", "http", "https"):
            return html.escape(u, quote=True)
        return "#"

    def _sanitize_img_src(url: str) -> str:
        if not url:
            return ""
        u = str(url).strip()
        if re.search(r"[\x00-\x1f\x7f\s\"'<>]", u):
            return ""
        try:
            parsed = urlparse(u)
        except Exception:
            return ""
        if parsed.scheme in ("http", "https"):
            return html.escape(u, quote=True)
        if parsed.scheme == "data" and u.lower().startswith("data:image/"):
            return html.escape(u, quote=True)
        return ""

    safe_company_name = html.escape(str(company_name), quote=True)
    safe_confidentiality = html.escape(str(confidentiality), quote=True)
    safe_report_title = html.escape(str(report.title or "Report"), quote=True)
    safe_company_logo_url = _sanitize_img_src(company_logo_url)
    primary_color = _sanitize_css_color(primary_color, "#1890ff")
    dark_color = _sanitize_css_color(dark_color, "#1a1a2e")

    report_type_raw = str(getattr(report, "report_type", "") or "").strip().lower()
    safe_report_type_class = report_type_raw if re.match(r"^[a-z0-9_-]{1,32}$", report_type_raw) else "executive"
    safe_report_type_label = html.escape(report_type_raw.replace("_", " ").title() or "Executive", quote=False)
    
    # Process markdown-style content to HTML
    def process_content(text: str) -> str:
        if not text:
            return ""
        
        lines = text.split('\n')
        html_lines = []
        in_list = False
        list_type = ''
        in_code_block = False
        
        for line in lines:
            trimmed = line.strip()
            
            # Code blocks
            if trimmed.startswith('```'):
                if in_code_block:
                    html_lines.append('</code></pre>')
                    in_code_block = False
                else:
                    if in_list:
                        html_lines.append('</ul>' if list_type == 'ul' else '</ol>')
                        in_list = False
                    html_lines.append('<pre class="code-block"><code>')
                    in_code_block = True
                continue
            
            if in_code_block:
                html_lines.append(html.escape(line, quote=False))
                continue
            
            # Horizontal rules
            if trimmed in ('---', '***', '___'):
                if in_list:
                    html_lines.append('</ul>' if list_type == 'ul' else '</ol>')
                    in_list = False
                html_lines.append('<hr />')
                continue
            
            # Headings
            if trimmed.startswith('#### '):
                if in_list:
                    html_lines.append('</ul>' if list_type == 'ul' else '</ol>')
                    in_list = False
                html_lines.append(f'<h4>{process_inline(trimmed[5:])}</h4>')
                continue
            if trimmed.startswith('### '):
                if in_list:
                    html_lines.append('</ul>' if list_type == 'ul' else '</ol>')
                    in_list = False
                html_lines.append(f'<h3>{process_inline(trimmed[4:])}</h3>')
                continue
            if trimmed.startswith('## '):
                if in_list:
                    html_lines.append('</ul>' if list_type == 'ul' else '</ol>')
                    in_list = False
                html_lines.append(f'<h2>{process_inline(trimmed[3:])}</h2>')
                continue
            if trimmed.startswith('# '):
                if in_list:
                    html_lines.append('</ul>' if list_type == 'ul' else '</ol>')
                    in_list = False
                html_lines.append(f'<h1>{process_inline(trimmed[2:])}</h1>')
                continue
            
            # Unordered lists
            if re.match(r'^[-*+]\s+', trimmed):
                if not in_list or list_type != 'ul':
                    if in_list:
                        html_lines.append('</ol>')
                    html_lines.append('<ul>')
                    in_list = True
                    list_type = 'ul'
                list_item = re.sub(r"^[-*+]\s+", "", trimmed)
                html_lines.append(f'<li>{process_inline(list_item)}</li>')
                continue
            
            # Ordered lists
            if re.match(r'^\d+\.\s+', trimmed):
                if not in_list or list_type != 'ol':
                    if in_list:
                        html_lines.append('</ul>')
                    html_lines.append('<ol>')
                    in_list = True
                    list_type = 'ol'
                list_item = re.sub(r"^\d+\.\s+", "", trimmed)
                html_lines.append(f'<li>{process_inline(list_item)}</li>')
                continue
            
            # Blockquotes
            if trimmed.startswith('> '):
                if in_list:
                    html_lines.append('</ul>' if list_type == 'ul' else '</ol>')
                    in_list = False
                html_lines.append(f'<blockquote>{process_inline(trimmed[2:])}</blockquote>')
                continue
            
            # Empty lines
            if trimmed == '':
                if in_list:
                    html_lines.append('</ul>' if list_type == 'ul' else '</ol>')
                    in_list = False
                continue
            
            # Regular paragraphs
            if in_list:
                html_lines.append('</ul>' if list_type == 'ul' else '</ol>')
                in_list = False
            html_lines.append(f'<p>{process_inline(trimmed)}</p>')
        
        if in_list:
            html_lines.append('</ul>' if list_type == 'ul' else '</ol>')
        
        return '\n'.join(html_lines)
    
    def process_inline(text: str) -> str:
        """Process inline markdown formatting with safe escaping and URL sanitization."""
        if not text:
            return ""

        link_tokens = {}
        token_idx = 0

        def _process_inline_without_links(raw: str) -> str:
            escaped = html.escape(raw, quote=False)
            escaped = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', escaped)
            escaped = re.sub(r'__([^_]+)__', r'<strong>\1</strong>', escaped)
            escaped = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', escaped)
            escaped = re.sub(r'_([^_]+)_', r'<em>\1</em>', escaped)
            escaped = re.sub(r'`([^`]+)`', r'<code>\1</code>', escaped)
            return escaped

        def _replace_link(match: re.Match) -> str:
            nonlocal token_idx
            label = match.group(1) or ""
            url = match.group(2) or ""
            # Use a token that won't be picked up by the simple emphasis regexes.
            token = f"@@PARSHULINKTOKEN{token_idx}@@"
            token_idx += 1
            link_tokens[token] = (label, url)
            return token

        tokenized = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', _replace_link, text)
        rendered = _process_inline_without_links(tokenized)

        for token, (label, url) in link_tokens.items():
            safe_label = _process_inline_without_links(label)
            safe_href = _sanitize_href(url)
            rendered = rendered.replace(
                token,
                f'<a href="{safe_href}" target="_blank" rel="noopener noreferrer">{safe_label}</a>',
            )

        return rendered

    sources_html_parts = []
    for a in articles:
        title = html.escape(str(a.title or ""), quote=False)
        href = _sanitize_href(getattr(a, "url", None) or "")
        if href != "#" and href:
            sources_html_parts.append(
                f'<div class="source-item"><a href="{href}" target="_blank" rel="noopener noreferrer">{title}</a></div>'
            )
        else:
            sources_html_parts.append(f"<div class=\"source-item\">{title}</div>")
    sources_html = "".join(sources_html_parts)

    company_logo_html = ""
    if safe_company_logo_url:
        company_logo_html = f'<img src="{safe_company_logo_url}" class="company-logo" alt="{safe_company_name}" />'
    
    # Build modern HTML document with branding
    html_content = f'''<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{safe_report_title} - {safe_company_name}</title>
  <style>
    :root {{
      --primary: {dark_color};
      --accent: {primary_color};
      --success: #10b981;
      --warning: #f59e0b;
      --danger: #ef4444;
      --gray-50: #f8fafc;
      --gray-100: #f1f5f9;
      --gray-200: #e2e8f0;
      --gray-400: #94a3b8;
      --gray-600: #475569;
      --gray-800: #1e293b;
    }}
    * {{ margin: 0; padding: 0; box-sizing: border-box; }}
    body {{
      font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
      line-height: 1.6;
      color: var(--gray-800);
      background: var(--gray-50);
      font-size: 14px;
    }}
    .classification-banner {{
      background: var(--danger);
      color: white;
      text-align: center;
      padding: 6px 16px;
      font-weight: 600;
      font-size: 11px;
      text-transform: uppercase;
      letter-spacing: 2px;
    }}
    .header {{
      background: var(--primary);
      color: white;
      padding: 24px 40px;
    }}
    .header-inner {{
      max-width: 900px;
      margin: 0 auto;
      display: flex;
      justify-content: space-between;
      align-items: center;
    }}
    .header-left {{
      display: flex;
      align-items: center;
      gap: 16px;
    }}
    .company-logo {{
      max-height: 40px;
      max-width: 140px;
      border-radius: 4px;
    }}
    .company-name {{
      font-size: 18px;
      font-weight: 600;
      letter-spacing: -0.025em;
    }}
    .header-right {{
      text-align: right;
      font-size: 13px;
      opacity: 0.9;
    }}
    .header-date {{
      font-weight: 500;
    }}
    .header-subtext {{
      font-size: 11px;
      opacity: 0.7;
      margin-top: 2px;
    }}
    main {{
      max-width: 900px;
      margin: 0 auto;
      padding: 32px 40px;
    }}
    .report-header {{
      margin-bottom: 32px;
      padding-bottom: 24px;
      border-bottom: 1px solid var(--gray-200);
    }}
    .report-title {{
      color: var(--primary);
      font-size: 28px;
      font-weight: 700;
      margin: 0 0 16px 0;
      letter-spacing: -0.025em;
    }}
    .report-meta {{
      display: flex;
      flex-wrap: wrap;
      gap: 20px;
      font-size: 13px;
      color: var(--gray-600);
    }}
    .meta-item {{
      display: flex;
      align-items: center;
      gap: 6px;
    }}
    .meta-label {{
      font-weight: 600;
      color: var(--gray-800);
    }}
    .report-type-badge {{
      display: inline-block;
      padding: 3px 10px;
      border-radius: 4px;
      font-size: 11px;
      font-weight: 600;
      text-transform: uppercase;
    }}
    .badge-executive {{ background: #dbeafe; color: #2563eb; }}
    .badge-technical {{ background: #d1fae5; color: #059669; }}
    .badge-comprehensive {{ background: #ede9fe; color: #7c3aed; }}
    .report-content {{
      background: white;
      border-radius: 8px;
      border: 1px solid var(--gray-200);
      padding: 32px;
      margin-bottom: 24px;
    }}
    h1 {{ 
      color: var(--primary); 
      font-size: 22px; 
      margin: 28px 0 14px 0; 
      font-weight: 600;
      letter-spacing: -0.025em;
    }}
    h2 {{ 
      color: var(--gray-800); 
      font-size: 18px; 
      margin: 24px 0 12px 0; 
      font-weight: 600;
      padding-bottom: 8px;
      border-bottom: 1px solid var(--gray-200);
    }}
    h3 {{ 
      color: var(--accent); 
      font-size: 15px; 
      margin: 20px 0 10px 0; 
      font-weight: 600;
    }}
    h4 {{ 
      color: var(--gray-600); 
      font-size: 14px; 
      margin: 16px 0 8px 0; 
      font-weight: 600;
    }}
    p {{ 
      margin: 0 0 12px 0; 
      color: var(--gray-600);
      line-height: 1.7;
    }}
    ul, ol {{ 
      margin: 12px 0; 
      padding-left: 24px; 
      color: var(--gray-600);
    }}
    li {{ 
      margin-bottom: 6px; 
      line-height: 1.6;
    }}
    blockquote {{
      border-left: 3px solid var(--accent);
      margin: 16px 0;
      padding: 12px 20px;
      background: var(--gray-50);
      color: var(--gray-600);
      font-style: italic;
      border-radius: 0 6px 6px 0;
    }}
    code {{
      background: var(--gray-100);
      padding: 2px 6px;
      border-radius: 4px;
      font-family: 'SF Mono', Monaco, Menlo, Consolas, monospace;
      font-size: 0.85em;
      color: #c026d3;
    }}
    .code-block {{
      background: var(--gray-800);
      color: #e2e8f0;
      padding: 16px 20px;
      border-radius: 8px;
      overflow-x: auto;
      margin: 16px 0;
      font-family: 'SF Mono', Monaco, Menlo, Consolas, monospace;
      font-size: 13px;
      line-height: 1.5;
    }}
    .code-block code {{
      background: transparent;
      color: inherit;
      padding: 0;
    }}
    hr {{ 
      border: none; 
      border-top: 1px solid var(--gray-200); 
      margin: 28px 0; 
    }}
    a {{ 
      color: var(--accent); 
      text-decoration: none; 
    }}
    a:hover {{ 
      text-decoration: underline; 
    }}
    strong {{ 
      color: var(--gray-800); 
      font-weight: 600;
    }}
    .sources-section {{
      background: white;
      border-radius: 8px;
      border: 1px solid var(--gray-200);
      padding: 24px 32px;
      margin-top: 24px;
    }}
    .sources-section h2 {{
      margin-top: 0;
      padding-bottom: 12px;
      border-bottom: 1px solid var(--gray-200);
    }}
    .source-item {{
      padding: 12px 16px;
      background: var(--gray-50);
      border-radius: 6px;
      margin-bottom: 8px;
      border-left: 3px solid var(--accent);
    }}
    .source-item a {{
      font-weight: 500;
      color: var(--gray-800);
    }}
    .source-item a:hover {{
      color: var(--accent);
    }}
    footer {{
      background: var(--primary);
      color: white;
      padding: 20px 40px;
      text-align: center;
      margin-top: 32px;
    }}
    .footer-classification {{
      font-size: 10px;
      opacity: 0.6;
      text-transform: uppercase;
      letter-spacing: 1px;
      margin-bottom: 6px;
    }}
    .footer-text {{
      font-size: 12px;
      opacity: 0.5;
    }}
    @media print {{
      body {{ background: white; }}
      main {{ padding: 20px; }}
      .report-content, .sources-section {{ border: 1px solid #ddd; box-shadow: none; }}
      .code-block {{ background: #f5f5f5; color: #333; border: 1px solid #ddd; }}
      .classification-banner, .header, footer {{ print-color-adjust: exact; -webkit-print-color-adjust: exact; }}
    }}
  </style>
</head>
<body>
  <div class="classification-banner">{safe_confidentiality}</div>
  
  <header class="header">
    <div class="header-inner">
      <div class="header-left">
        {company_logo_html}
        <div class="company-name">{safe_company_name}</div>
      </div>
      <div class="header-right">
        <div class="header-date">{report.generated_at.strftime('%B %d, %Y')}</div>
        <div class="header-subtext">Threat Intelligence Report</div>
      </div>
    </div>
  </header>

  <main>
    <div class="report-header">
      <h1 class="report-title">{safe_report_title}</h1>
      <div class="report-meta">
        <div class="meta-item">
           <span class="meta-label">Type:</span>
           <span class="report-type-badge badge-{safe_report_type_class}">{safe_report_type_label}</span>
         </div>
        <div class="meta-item">
          <span class="meta-label">Generated:</span>
          <span>{report.generated_at.strftime('%B %d, %Y at %H:%M UTC')}</span>
        </div>
        <div class="meta-item">
          <span class="meta-label">Articles:</span>
          <span>{len(articles)}</span>
        </div>
        <div class="meta-item">
          <span class="meta-label">Classification:</span>
          <span style="color: var(--danger); font-weight: 600;">{html.escape((confidentiality.split('-')[0].strip() if '-' in confidentiality else confidentiality), quote=False)}</span>
        </div>
      </div>
    </div>
    
    <section class="report-content">
      {process_content(report.content)}
    </section>
    
    <section class="sources-section">
      <h2>Source Articles</h2>
      {sources_html}
    </section>
  </main>
  
  <footer>
    <div class="footer-classification">{safe_confidentiality}</div>
    <div class="footer-text">{safe_company_name} • Threat Intelligence Report • Generated by Parshu</div>
  </footer>
</body>
</html>'''
    
    logger.info("html_report_exported", report_id=report_id, user_id=current_user.id)
    
    return StreamingResponse(
        iter([html_content]),
        media_type="text/html",
        headers={"Content-Disposition": f"inline; filename=report_{report_id}.html"}
    )
